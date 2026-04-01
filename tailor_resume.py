#!/usr/bin/env python3
"""
Resume Tailor - ATS optimizer using Claude chat (no API key required)

Workflow:
  Step 1:  python tailor_resume.py --mode prep   [--jd jd.txt]
           → writes prompt.txt  (paste into claude.ai, copy the JSON reply)
           → save Claude's reply as response.json

  Step 2:  python tailor_resume.py --mode apply
           → reads response.json, writes tailored .docx + .pdf
"""

import os
import sys
import json
import shutil
import subprocess
import argparse
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# ─── CONFIGURATION ─────────────────────────────────────────────────────────────
TEMPLATE_CV  = Path("AamirAli_Resume_ATS.docx")
OUTPUT_DIR   = Path("generated_resumes")
PROMPT_FILE  = Path("prompt.txt")
RESPONSE_FILE = Path("response.json")

# Namespace shortcut
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# Current CV content — used as context for Claude
CV_CONTENT = """
=== SKILLS ===
• Programming Languages: Python, C++
• Databases: CosmosDB, MSSQL, NoSQL, SnowSQL, kdb+
• Cloud Platforms: AWS (S3, EC2, EMR, Glue, Redshift), Snowflake, Azure (Data Factory, Databricks, Synapse), GCP, Terraform
• Data Engineering Tools: DBT, PySpark, Kafka, Data Vault 2.0, CDC, ELT, Looker, Looker API, Apache Airflow (DAGs), Fivetran, MLOps, Model Deployment Pipelines, Model Registry, Experiment Tracking, Model Versioning
• CI/CD: Git, CI/CD pipelines, Azure Function, Agile/Scrum, Automated Testing, L3 Support
• Data techniques: ETL, ELT, Data Modelling, Physical Modelling, Data Ingestion, Data Migration, Streaming, Data Warehousing, Data Analysis, CDC, Reconciliation, Semantic Layers, Rules Engines, Observability, Self-Service Analytics, Near-Real-Time Inference, Feature Drift Monitoring, Model Rollback

=== EXPERIENCE ===

[1] Senior Data Engineer | Build A Rocket Boy | London | 05/2025 - Current
• Designed and developed DBT data models with semantic layers integrating 100+ source systems into a centralized Snowflake warehouse processing 120+ TB of structured and semi-structured data.
• Led end-to-end migration from dbt Core CLI to dbt Fusion, independently owning the transition and establishing Snowflake target schemas, ELT pipelines, and CI/CD pipelines with automated testing.
• Automated ELT workflows and near-real-time inference pipelines using Fivetran connectors and Python on Terraform-provisioned AWS, with feature drift monitoring and observability frameworks ensuring model and pipeline reliability.
• Designed and deployed Gen AI model pipelines using MLOps best practices, establishing model deployment pipelines, model registry, experiment tracking, model versioning, and rollback procedures for production AI workloads over Snowflake.

[2] Lead Data Engineer | Castleton Commodities Int. | London | 08/2024 - 04/2025
• Architected and deployed scalable ELT pipelines in Snowflake, DBT, and AWS/Azure with CDC-based ingestion and Apache Airflow DAG orchestration, supporting Natural Gas and Power trading desks, settlements, and risk functions.
• Developed near-real-time inference data pipelines and reconciliation transformations for pre-trade and market data (prices, positions, exposures, curves) supporting P&L reporting, mark-to-market, and VaR analysis.
• Built rules-engine-driven data quality checks with feature drift monitoring and Power BI dashboards for front-office and risk stakeholders, delivering transparency into P&L, working capital, and trade settlements.
• Provided L3 support and authored technical documentation for all data warehouse changes; developed hotfixes for production incidents and mentored junior developers on data engineering best practices.
• Integrated regulatory reporting datasets (REMIT, EMIR) into governed data models ensuring full lineage, auditability, and compliance using Terraform-managed infrastructure.

[3] Lead Data Engineer | Sodexo | London | 05/2023 - 01/2024
• Led CDC-based data ingestion for the AMETA region using Azure ADF, Databricks, PySpark, and Fivetran-style connectors, building observability frameworks that eliminated third-party tools and reduced external app costs by 80%.
• Reported directly to CTO as Engineering lead for designing and developing the fintech transaction self-service analytics platform, providing reconciliation and rich visual summary of 1M+ daily user decisions.
• Designed Agile/Scrum-driven data pipelines using Azure Data Factory with an Orchestration framework and Azure Functions, delivering self-service analytics with observability controls and L3 support for the data warehouse.
• Mentored junior developers and led Data Vault 2.0 physical modelling implementation with rules engines for data quality, reducing data redundancy and accelerating SDLC delivery cycles.

[4] Data Engineer | Wherescape USA | (US Remote) | 07/2019 - 05/2023
• BAT (British American Tobacco) - Implemented Data Migration scripts and Azure Synapse Migration with automated testing suites. Implemented multi-cluster Snowflake warehouses with workload isolation ensuring SLA compliance and L3 support.
• Designed Snowflake external tables on S3 for semi-structured data (JSON, Parquet) using physical modelling best practices, integrating IoT/supply-chain feeds via ELT into core analytics pipelines.
• AMD Malaysia - Migrated data warehouse from Informatica to Wherescape using Agile/Scrum methodology. Created SQL transformations from Oracle to SQL Server, built incremental ELT pipeline loading 1B records with 80% efficiency gain.
"""

SYSTEM_PROMPT = """You are an expert ATS resume optimizer and professional resume writer specializing in data engineering roles.
Your job is to tailor resumes to specific job descriptions to maximize ATS keyword scores.

Rules for rewriting bullet points:
- Each bullet must be approximately 180-220 characters (about 2 printed lines)
- Start every bullet with a strong past-tense action verb
- Embed keywords from the JD naturally into existing experience facts
- Add realistic metrics and quantified results where missing
- Use STAR format: action + context + measurable result
- Do NOT invent companies, technologies the candidate never used, or impossible claims
- DO strategically embellish scale, scope and impact to sound more impressive
- Wrap JD keywords and technologies in **double asterisks** so they render as bold in the output document. Example: "Engineered **ELT** pipelines in **Snowflake** processing 5TB daily"

For skills: only add technologies that are adjacent to candidate's existing stack."""


# ─── PROMPT BUILDER ────────────────────────────────────────────────────────────

def build_prompt(jd_text: str) -> str:
    """Build the full prompt to paste into Claude chat."""
    return f"""{SYSTEM_PROMPT}

---

Analyze this job description and tailor the resume to maximize ATS score.

JOB DESCRIPTION:
{jd_text}

CURRENT RESUME:
{CV_CONTENT}

Return ONLY a valid JSON object with this exact structure (no markdown, no explanation):
{{
  "company_name": "Company name from JD (short form, e.g. 'Google')",
  "role_title": "Exact job title from JD",
  "ats_match_percentage": <integer 0-100>,
  "missing_keywords": ["keyword1", "keyword2", "keyword3"],
  "selection_likelihood": "Low/Medium/High - one sentence reason",
  "skills": {{
    "Programming Languages": "Python, C++, [add from JD if relevant]",
    "Databases": "CosmosDB, MSSQL, NoSQL, SnowSQL, kdb+, [add from JD]",
    "Cloud Platforms": "AWS (S3, EC2, EMR, Glue, Redshift), Snowflake, Azure (Data Factory, Databricks, Synapse), GCP, Terraform, [add from JD]",
    "Data Engineering Tools": "DBT, PySpark, Kafka, Data Vault 2.0, CDC, ELT, Looker, Looker API, Apache Airflow (DAGs), Fivetran, MLOps, Model Deployment Pipelines, Model Registry, Experiment Tracking, Model Versioning, [add from JD]",
    "CI/CD": "Git, CI/CD pipelines, Azure Function, Agile/Scrum, Automated Testing, L3 Support, [add from JD]",
    "Data techniques": "ETL, ELT, Data Modelling, Physical Modelling, Data Ingestion, Data Migration, Streaming, Data Warehousing, Data Analysis, CDC, Reconciliation, Semantic Layers, Rules Engines, Observability, Self-Service Analytics, Near-Real-Time Inference, Feature Drift Monitoring, Model Rollback, [add from JD]"
  }},
  "experience": {{
    "Build A Rocket Boy": [
      "Bullet 1 — ~200 chars, past tense verb, JD keywords, metric",
      "Bullet 2",
      "Bullet 3",
      "Bullet 4"
    ],
    "Castleton Commodities Int.": [
      "Bullet 1", "Bullet 2", "Bullet 3", "Bullet 4", "Bullet 5"
    ],
    "Sodexo": [
      "Bullet 1", "Bullet 2", "Bullet 3", "Bullet 4"
    ],
    "Wherescape USA": [
      "Bullet 1", "Bullet 2", "Bullet 3"
    ]
  }}
}}"""


# ─── XML HELPERS ───────────────────────────────────────────────────────────────

def wtag(name: str) -> str:
    return f"{{{W}}}{name}"


def get_para_full_text(para) -> str:
    """Get all text from a paragraph including text from all runs."""
    return "".join(
        (t.text or "") for r in para._p.findall(wtag("r"))
        for t in [r.find(wtag("t"))] if t is not None
    )


def is_bullet_para(para) -> bool:
    """True if paragraph starts with a known bullet character."""
    text = get_para_full_text(para).strip()
    return bool(text) and (text[0] in "◆◆◆•-")


def replace_skill_values(para, category: str, new_values: str):
    """
    In a skill paragraph like:  ◆ [bold]Category[/bold]: val1, val2
    replace everything after the category run with ': new_values'.
    """
    p = para._p
    children = list(p)

    cat_pos = -1
    for i, child in enumerate(children):
        if child.tag == wtag("r"):
            t = child.find(wtag("t"))
            if t is not None and category in (t.text or ""):
                cat_pos = i
                break

    if cat_pos == -1:
        return

    sample_rpr = None
    to_remove = []

    for child in children[cat_pos + 1:]:
        if child.tag == wtag("r"):
            if child.find(wtag("br")) is not None:
                break
            t = child.find(wtag("t"))
            if t is not None:
                t_text = t.text or ""
                if t_text.strip() and t_text.strip()[0] in "◆◆◆•-":
                    break
                if sample_rpr is None and t_text.strip() and t_text.strip() != ":":
                    rpr = child.find(wtag("rPr"))
                    if rpr is not None:
                        sample_rpr = deepcopy(rpr)
            to_remove.append(child)
        elif child.tag == wtag("proofErr"):
            to_remove.append(child)

    for r in to_remove:
        p.remove(r)

    updated_children = list(p)
    insert_before_idx = None
    passed_cat = False
    for i, child in enumerate(updated_children):
        if child.tag == wtag("r"):
            t = child.find(wtag("t"))
            if t is not None and category in (t.text or ""):
                passed_cat = True
            if passed_cat and child.find(wtag("br")) is not None:
                insert_before_idx = i
                break

    new_run = etree.Element(wtag("r"))
    if sample_rpr is not None:
        new_run.append(deepcopy(sample_rpr))
    new_t = etree.SubElement(new_run, wtag("t"))
    new_t.text = f": {new_values}"
    new_t.set(XML_SPACE, "preserve")

    if insert_before_idx is not None:
        p.insert(insert_before_idx, new_run)
    else:
        p.append(new_run)


def replace_bullet_content(para, new_text: str):
    """
    Replace text in a bullet paragraph with new_text.
    - Preserves bullet character run and paragraph spacing (pPr)
    - Normalises text runs: clean font+size, black color, no stray bold/italic
    - Parses **bold** markers in new_text and renders them as bold runs
    """
    import re as _re
    p = para._p
    children = list(p)

    bullet_run = None
    font_name  = "Calibri"
    size_val   = None

    for child in children:
        if child.tag != wtag("r"):
            continue
        t = child.find(wtag("t"))
        if t is None:
            continue
        stripped = (t.text or "").strip()
        if stripped and len(stripped) <= 2 and stripped[0] in "◆◆◆•-":
            bullet_run = child
        elif stripped and size_val is None:
            # Extract font name and size only — ignore bold/italic/color
            rpr = child.find(wtag("rPr"))
            if rpr is not None:
                rf = rpr.find(wtag("rFonts"))
                if rf is not None:
                    font_name = rf.get(wtag("ascii"), font_name)
                sz = rpr.find(wtag("sz"))
                if sz is not None:
                    size_val = sz.get(wtag("val"))

    # Remove all runs (keep pPr)
    for child in list(p):
        if child.tag != wtag("pPr"):
            p.remove(child)

    # Re-add bullet run unchanged
    if bullet_run is not None:
        p.append(deepcopy(bullet_run))

    def make_rpr(bold: bool = False) -> etree._Element:
        """Build a clean rPr: font, size, black color. Bold only if requested."""
        rpr = etree.Element(wtag("rPr"))
        rf = etree.SubElement(rpr, wtag("rFonts"))
        rf.set(wtag("ascii"), font_name)
        rf.set(wtag("hAnsi"), font_name)
        rf.set(wtag("cs"),    font_name)
        if bold:
            etree.SubElement(rpr, wtag("b"))
            etree.SubElement(rpr, wtag("bCs"))
        if size_val:
            etree.SubElement(rpr, wtag("sz")).set(wtag("val"), size_val)
            etree.SubElement(rpr, wtag("szCs")).set(wtag("val"), size_val)
        # Force text color to auto (black) — prevents blue bleed from bullet run
        etree.SubElement(rpr, wtag("color")).set(wtag("val"), "auto")
        return rpr

    def add_run(text: str, bold: bool = False):
        r = etree.SubElement(p, wtag("r"))
        r.append(make_rpr(bold))
        t = etree.SubElement(r, wtag("t"))
        t.text = text
        t.set(XML_SPACE, "preserve")

    # Space after bullet
    add_run(" ")

    # Split on **...** markers and emit normal / bold runs alternately
    parts = _re.split(r'\*\*(.+?)\*\*', new_text)
    for i, part in enumerate(parts):
        if part:
            add_run(part, bold=(i % 2 == 1))


# ─── DOCUMENT UPDATE ──────────────────────────────────────────────────────────

SKILL_CATEGORIES = [
    "Programming Languages",
    "Databases",
    "Cloud Platforms",
    "Data Engineering Tools",
    "CI/CD",
    "Data techniques",
]

COMPANIES = [
    "Build A Rocket Boy",
    "Castleton Commodities Int.",
    "Sodexo",
    "Wherescape USA",
]


def update_document(doc: Document, data: dict) -> Document:
    """Apply Claude's tailored content to the document."""
    skills_data = data.get("skills", {})
    exp_data = data.get("experience", {})

    current_company = None
    bullet_counters = {c: 0 for c in COMPANIES}

    for para in doc.paragraphs:
        text = get_para_full_text(para)

        for cat in SKILL_CATEGORIES:
            if cat in text and cat in skills_data and skills_data[cat]:
                replace_skill_values(para, cat, skills_data[cat])
                break

        for company in COMPANIES:
            if company in text and not is_bullet_para(para):
                current_company = company
                break

        if current_company and is_bullet_para(para):
            bullets = exp_data.get(current_company, [])
            idx = bullet_counters[current_company]
            if idx < len(bullets):
                replace_bullet_content(para, bullets[idx])
                bullet_counters[current_company] = idx + 1

    return doc


# ─── PDF CONVERSION ───────────────────────────────────────────────────────────

def convert_to_pdf(docx_path: Path) -> Path | None:
    """Convert docx → PDF. Tries docx2pdf (requires MS Word) then LibreOffice."""
    pdf_path = docx_path.with_suffix(".pdf")

    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except ImportError:
        pass
    except Exception as e:
        print(f"  docx2pdf error: {e}")

    for soffice in ["soffice", "libreoffice"]:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                return pdf_path
        except (FileNotFoundError, subprocess.SubprocessError):
            continue

    print("  PDF skipped - install docx2pdf ('pip install docx2pdf') or LibreOffice")
    return None


# ─── INPUT ────────────────────────────────────────────────────────────────────

def get_jd_text(jd_file: str | None) -> str:
    """Return job description text from file, jd.txt, or interactive paste."""
    if jd_file:
        path = Path(jd_file)
        if not path.exists():
            print(f"❌  File not found: {jd_file}")
            sys.exit(1)
        return path.read_text(encoding="utf-8")

    auto = Path("jd.txt")
    if auto.exists():
        print("  Found jd.txt - using it as job description.")
        return auto.read_text(encoding="utf-8")

    print("\nPaste the job description below.")
    print("When finished, press Enter, type END on a new line, then press Enter:")
    print("-" * 60)
    lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip().upper() == "END":
            break
        lines.append(line)
    return "\n".join(lines)


def try_copy_to_clipboard(text: str) -> bool:
    """Try to copy text to clipboard. Returns True on success."""
    try:
        import subprocess
        proc = subprocess.run(
            ["clip"], input=text.encode("utf-8"),
            capture_output=True
        )
        return proc.returncode == 0
    except Exception:
        return False


# ─── ATS COMPLIANCE CHECKER ──────────────────────────────────────────────────

def check_ats_compliance(doc_path: Path) -> list:
    """
    Inspect the .docx structure for ATS-hostile formatting.
    Returns a list of dicts: {level: PASS|WARN|FAIL, check: str, detail: str}
    """
    import re as _re
    findings = []
    doc = Document(str(doc_path))
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # 1. Tables
    if doc.tables:
        findings.append({
            "level": "FAIL",
            "check": "Tables detected",
            "detail": f"{len(doc.tables)} table(s) found — ATS parsers often skip content inside tables entirely.",
        })
    else:
        findings.append({"level": "PASS", "check": "No tables", "detail": ""})

    # 2. Text boxes
    body_xml = doc.element.xml
    if "txbxContent" in body_xml or "w:txbx" in body_xml:
        findings.append({
            "level": "FAIL",
            "check": "Text boxes detected",
            "detail": "Content in text boxes is invisible to most ATS parsers — move to plain paragraphs.",
        })
    else:
        findings.append({"level": "PASS", "check": "No text boxes", "detail": ""})

    # 3. Multi-column layout
    for section in doc.sections:
        cols = section._sectPr.find(f"{{{W_NS}}}cols")
        if cols is not None:
            num = cols.get(f"{{{W_NS}}}num")
            if num and int(num) > 1:
                findings.append({
                    "level": "FAIL",
                    "check": f"Multi-column layout ({num} columns)",
                    "detail": "Column layouts confuse ATS parsers — use a single-column format.",
                })
                break

    # 4. Headers / footers with content
    for section in doc.sections:
        for label, hf in [("Header", section.header), ("Footer", section.footer)]:
            if hf is None:
                continue
            text = " ".join(p.text.strip() for p in hf.paragraphs).strip()
            if text:
                findings.append({
                    "level": "WARN",
                    "check": f"{label} contains text",
                    "detail": f'"{text[:80]}" — ATS may ignore header/footer content.',
                })

    # 5. Images / drawings
    ns_draw = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    drawings = doc.element.findall(f".//{{{W_NS}}}drawing")
    if drawings or doc.inline_shapes:
        count = len(drawings) or len(doc.inline_shapes)
        findings.append({
            "level": "WARN",
            "check": f"Images or graphics detected ({count})",
            "detail": "ATS cannot read text in images. Remove if any critical info is embedded.",
        })
    else:
        findings.append({"level": "PASS", "check": "No images or graphics", "detail": ""})

    # 6. Standard section headings present
    all_text_upper = "\n".join(p.text.strip().upper() for p in doc.paragraphs)
    for section_name in ["EXPERIENCE", "SKILLS", "EDUCATION"]:
        if section_name not in all_text_upper:
            findings.append({
                "level": "WARN",
                "check": f'Section heading "{section_name}" not found',
                "detail": "ATS uses standard headings to categorise resume content.",
            })

    # 7. Font check
    SAFE_FONTS = {
        "arial", "calibri", "calibri (body)", "times new roman", "helvetica",
        "georgia", "garamond", "cambria", "trebuchet ms", "verdana", "tahoma",
    }
    unusual = set()
    for para in doc.paragraphs:
        for run in para.runs:
            name = (run.font.name or "").strip().lower()
            if name and name not in SAFE_FONTS:
                unusual.add(run.font.name.strip())
    if unusual:
        findings.append({
            "level": "WARN",
            "check": "Non-standard fonts detected",
            "detail": f"{', '.join(sorted(unusual))} — prefer Arial, Calibri, or Times New Roman.",
        })
    else:
        findings.append({"level": "PASS", "check": "Standard fonts used", "detail": ""})

    # 8. Contact info in body (not locked away in header)
    body_text = " ".join(p.text for p in doc.paragraphs)
    if not _re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", body_text):
        findings.append({
            "level": "WARN",
            "check": "Email address not found in document body",
            "detail": "Place your email in the main body so ATS can reliably extract it.",
        })
    else:
        findings.append({"level": "PASS", "check": "Email found in document body", "detail": ""})

    if not _re.search(r"[\+\d][\d\s\-\(\)]{7,}", body_text):
        findings.append({
            "level": "WARN",
            "check": "Phone number not found in document body",
            "detail": "Place your phone number in the main body so ATS can reliably extract it.",
        })
    else:
        findings.append({"level": "PASS", "check": "Phone number found in document body", "detail": ""})

    # 9. Excessive use of special bullet characters
    special_bullet_paras = sum(
        1 for p in doc.paragraphs
        if p.text.strip() and p.text.strip()[0] in "\u25c6\u2022\u25a0\u25ba\u2013\u2014"
    )
    total_paras = sum(1 for p in doc.paragraphs if p.text.strip())
    if total_paras and (special_bullet_paras / total_paras) > 0.6:
        findings.append({
            "level": "WARN",
            "check": "Heavy use of special bullet characters",
            "detail": "Some ATS parsers strip special chars — standard hyphens or plain bullets are safer.",
        })

    return findings


def mode_check(doc_file: str | None):
    """Check a .docx file for ATS compliance issues."""
    if doc_file:
        target = Path(doc_file)
    else:
        target = TEMPLATE_CV

    if not target.exists():
        print(f"[ERR]  File not found: {target}")
        print(f"       Usage: python tailor_resume.py --mode check --doc <file.docx>")
        sys.exit(1)

    print(f"  Checking: {target.name}\n")

    findings = check_ats_compliance(target)

    fails  = [f for f in findings if f["level"] == "FAIL"]
    warns  = [f for f in findings if f["level"] == "WARN"]
    passes = [f for f in findings if f["level"] == "PASS"]

    # Summary line
    score_label = "GOOD" if not fails and len(warns) <= 2 else ("NEEDS WORK" if not fails else "AT RISK")
    print(f"  Result : {score_label}  |  {len(fails)} fail(s)  {len(warns)} warning(s)  {len(passes)} pass(es)")
    print("-" * 60)

    for f in fails:
        print(f"  [FAIL]  {f['check']}")
        if f["detail"]:
            print(f"          {f['detail']}")

    for f in warns:
        print(f"  [WARN]  {f['check']}")
        if f["detail"]:
            print(f"          {f['detail']}")

    for f in passes:
        print(f"  [PASS]  {f['check']}")

    print("-" * 60)
    if fails:
        print("\n  Action: fix FAIL items before submitting — they can cause ATS to drop your resume.")
    elif warns:
        print("\n  Action: review warnings and fix where possible.")
    else:
        print("\n  Your resume looks ATS-friendly!")
    print()


# ─── MODES ────────────────────────────────────────────────────────────────────

def mode_prep(jd_file: str | None):
    """Step 1: build prompt and save to prompt.txt."""
    print("  [Step 1 of 2]  Generating prompt...\n")

    jd_text = get_jd_text(jd_file)
    if not jd_text.strip():
        print("❌  No job description provided.")
        sys.exit(1)

    prompt = build_prompt(jd_text)
    PROMPT_FILE.write_text(prompt, encoding="utf-8")

    copied = try_copy_to_clipboard(prompt)

    print(f"  [OK]  Prompt written to:  {PROMPT_FILE.resolve()}")
    if copied:
        print("  [OK]  Prompt also copied to clipboard.")
    print()
    print("  Next steps:")
    print("  " + "-" * 53)
    if not copied:
        print(f"  1. Open {PROMPT_FILE} and copy its entire contents.")
    else:
        print("  1. The prompt is already in your clipboard.")
    print("  2. Paste into claude.ai (or any Claude chat window).")
    print("  3. Copy Claude's entire JSON reply.")
    print(f"  4. Save it as:  {RESPONSE_FILE.resolve()}")
    print(f"  5. Run:  python tailor_resume.py --mode apply")
    print("  " + "-" * 53)


def mode_apply():
    """Step 2: read response.json and produce the tailored docx + pdf."""
    print("  [Step 2 of 2]  Applying Claude's response...\n")

    if not RESPONSE_FILE.exists():
        print(f"❌  {RESPONSE_FILE} not found.")
        print(f"    Run --mode prep first, then save Claude's reply as {RESPONSE_FILE}.")
        sys.exit(1)

    if not TEMPLATE_CV.exists():
        print(f"❌  Template not found: {TEMPLATE_CV}")
        sys.exit(1)

    raw = RESPONSE_FILE.read_text(encoding="utf-8").strip()

    # Strip markdown code fences if Claude added them
    if "```json" in raw:
        raw = raw.split("```json")[1].split("```")[0].strip()
    elif "```" in raw:
        raw = raw.split("```")[1].split("```")[0].strip()

    # Remove invisible control characters that browsers sometimes inject on copy
    import re
    raw = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', raw)

    try:
        # strict=False allows literal control chars inside strings (e.g. stray \r)
        data = json.loads(raw, strict=False)
    except json.JSONDecodeError as e:
        print(f"[ERR]  Could not parse {RESPONSE_FILE} as JSON: {e}")
        print("       Make sure you saved only Claude's JSON reply (no extra text).")
        sys.exit(1)

    company    = data.get("company_name", "Company")
    role       = data.get("role_title", "")
    ats_pct    = data.get("ats_match_percentage", 0)
    missing    = data.get("missing_keywords", [])
    likelihood = data.get("selection_likelihood", "N/A")

    print("-" * 60)
    print(f"  Company          : {company}")
    print(f"  Role             : {role}")
    print(f"  ATS Match Score  : {ats_pct}%")
    print(f"  Selection Chance : {likelihood}")
    if missing:
        shown  = missing[:10]
        extra  = len(missing) - 10
        kw_str = ", ".join(shown)
        if extra > 0:
            kw_str += f" (+{extra} more)"
        print(f"  Missing Keywords : {kw_str}")
    print("-" * 60 + "\n")

    print("  Tailoring resume...")
    doc = Document(str(TEMPLATE_CV))
    doc = update_document(doc, data)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = company.replace(" ", "_").replace("/", "-").replace("\\", "-")
    out_docx  = OUTPUT_DIR / f"{safe_name}.docx"
    doc.save(str(out_docx))
    print(f"  [OK]  Word saved  : {out_docx.name}")

    print("  Converting to PDF...")
    pdf_result = convert_to_pdf(out_docx)
    if pdf_result:
        print(f"  [OK]  PDF saved   : {pdf_result.name}")

    print(f"\nDone!  Files are in: {OUTPUT_DIR.resolve()}\n")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="ATS Resume Tailor - manual Claude chat workflow",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Workflow:
  Step 1:  python tailor_resume.py --mode prep   [--jd jd.txt]
           Writes prompt.txt - paste into claude.ai, save reply as response.json

  Step 2:  python tailor_resume.py --mode apply
           Reads response.json, writes tailored .docx + .pdf
""",
    )
    parser.add_argument(
        "--mode", choices=["prep", "apply", "check"], required=True,
        help="prep = generate prompt | apply = apply response | check = ATS compliance check"
    )
    parser.add_argument("--jd", metavar="FILE", help="Path to job description .txt file (prep mode only)")
    parser.add_argument("--doc", metavar="FILE", help="Path to .docx to check (check mode, default: template CV)")
    args = parser.parse_args()

    print("=" * 60)
    print("   RESUME TAILOR  -  ATS Optimizer")
    print("=" * 60)
    print()

    if args.mode == "prep":
        mode_prep(args.jd)
    elif args.mode == "apply":
        mode_apply()
    else:
        mode_check(args.doc)


if __name__ == "__main__":
    main()
